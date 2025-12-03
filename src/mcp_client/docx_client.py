from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from PIL import Image
from io import BytesIO
import os
from typing import Optional, List, Dict, Any
from src.utils.logger import LOGGER

class LocalDocxGenerator:
    """
    Local DOCX document generator using python-docx.
    """

    def __init__(self, output_dir: str = "./.results"):
        """
        Initialize local DOCX generator.

        Args:
            output_dir: Directory to save generated documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        LOGGER.info(f"LocalDocxGenerator initialized with output_dir: {output_dir}")

    def create_document(
        self,
        title: str,
        content_blocks: List[Dict[str, Any]],
        output_filename: Optional[str] = None
    ) -> str:
        """
        Create a DOCX document with formatted content and images.

        Args:
            title: Document title
            content_blocks: List of content blocks (text, images, etc.)
            output_filename: Optional custom filename

        Returns:
            Path to the created document
        """
        LOGGER.info("=== DOCX CREATION START ===")
        LOGGER.info(f"Title: {title}")
        LOGGER.info(f"Content blocks: {len(content_blocks)}")

        try:
            # Create document
            doc = Document()

            # Add title
            LOGGER.info(f"Adding title: {title}")
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Process content blocks
            LOGGER.info("Processing content blocks...")
            for i, block in enumerate(content_blocks):
                block_type = block.get("type")
                LOGGER.debug(f"Block {i+1}/{len(content_blocks)}: type={block_type}")

                if block_type == "heading":
                    level = block.get("level", 1)
                    _ = doc.add_heading(block.get("text", ""), level)

                elif block_type == "paragraph":
                    para = doc.add_paragraph(block.get("text", ""))
                    # Apply formatting if specified
                    if block.get("bold"):
                        para.runs[0].bold = True
                    if block.get("italic"):
                        para.runs[0].italic = True

                elif block_type == "image":
                    image_url = block.get("url")
                    caption = block.get("caption", "")

                    LOGGER.info(f"Processing image: {image_url[:100]}...")

                    try:
                        # Download image
                        LOGGER.info(f"Downloading image from: {image_url}")
                        response = requests.get(image_url, timeout=30)

                        if response.status_code == 200:
                            LOGGER.info(f"Image downloaded successfully ({len(response.content)} bytes)")

                            # Load image to check dimensions
                            img = Image.open(BytesIO(response.content))
                            LOGGER.info(f"Image opened: {img.size}, format: {img.format}")

                            # Save to temp file
                            temp_path = os.path.join(self.output_dir, f"temp_{block.get('id', 'img')}.jpg")
                            LOGGER.info(f"Saving to temp file: {temp_path}")

                            img.save(temp_path, "JPEG")
                            LOGGER.info("Temp file saved successfully")

                            # Add to document (max width 6 inches)
                            LOGGER.info("Adding picture to document...")
                            doc.add_picture(temp_path, width=Inches(6))
                            LOGGER.info("Picture added successfully")

                            # Add caption
                            if caption:
                                caption_para = doc.add_paragraph(caption)
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_para.runs[0].italic = True
                                caption_para.runs[0].font.size = Pt(9)

                            # Clean up temp file
                            os.remove(temp_path)
                            LOGGER.info("Temp file removed")
                        else:
                            LOGGER.warning(f"Failed to download image: HTTP {response.status_code}")
                            doc.add_paragraph(f"[Imagem não disponível: {caption or image_url}]")

                    except Exception as e:
                        LOGGER.error(f"Error adding image: {e}", exc_info=True)
                        doc.add_paragraph(f"[Imagem: {caption or image_url}]")

                elif block_type == "bullet_list":
                    items = block.get("items", [])
                    for item in items:
                        doc.add_paragraph(item, style='List Bullet')

            # Save document
            LOGGER.info("Saving document...")
            if not output_filename:
                import re
                safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
                output_filename = f"{safe_title}.docx"
                LOGGER.info(f"Generated filename: {output_filename}")

            output_path = os.path.join(self.output_dir, output_filename)
            LOGGER.info(f"Full output path: {output_path}")

            doc.save(output_path)
            LOGGER.info(f"Document saved successfully to: {output_path}")

            # Verify file exists
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                LOGGER.info(f"File verified: {file_size} bytes")
            else:
                LOGGER.error(f"File was not created at: {output_path}")

            LOGGER.info(f"Document saved: {output_path}")
            return output_path

        except Exception as e:
            LOGGER.error(f"Error creating DOCX: {e}", exc_info=True)
            return ""