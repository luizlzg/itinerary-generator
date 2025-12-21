from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from PIL import Image
from io import BytesIO
import os
from typing import Optional, List, Dict, Any
from src.utils.utilities import plot_clusters_on_basemap
from src.utils.logger import LOGGER


# Modern color palette
COLORS = {
    "primary": RGBColor(0x00, 0x7A, 0xCC),      # Modern blue
    "secondary": RGBColor(0x00, 0xB4, 0xD8),    # Light blue/teal
    "accent": RGBColor(0xFF, 0x69, 0x00),       # Orange accent
    "dark": RGBColor(0x2D, 0x3A, 0x4A),         # Dark gray-blue
    "text": RGBColor(0x33, 0x33, 0x33),         # Dark text
    "light_text": RGBColor(0x66, 0x66, 0x66),   # Light gray text
    "background": RGBColor(0xF8, 0xF9, 0xFA),   # Light background
}

# Language-specific labels for document generation
DOCX_LABELS = {
    "en": {
        "travel_itinerary": "Travel Itinerary",
        "route_map": "Route Map",
        "map_legend": "Colored dots indicate attractions grouped by day",
        "image_not_available": "Image not available",
        "map_not_available": "Map not available",
        "day_prefix": "Day",
    },
    "pt-br": {
        "travel_itinerary": "Roteiro de Viagem",
        "route_map": "Mapa do Roteiro",
        "map_legend": "Pontos coloridos indicam atrações agrupadas por dia",
        "image_not_available": "Imagem não disponível",
        "map_not_available": "Mapa não disponível",
        "day_prefix": "Dia",
    },
    "es": {
        "travel_itinerary": "Itinerario de Viaje",
        "route_map": "Mapa de la Ruta",
        "map_legend": "Los puntos de colores indican atracciones agrupadas por día",
        "image_not_available": "Imagen no disponible",
        "map_not_available": "Mapa no disponible",
        "day_prefix": "Día",
    },
    "fr": {
        "travel_itinerary": "Itinéraire de Voyage",
        "route_map": "Carte de l'Itinéraire",
        "map_legend": "Les points colorés indiquent les attractions regroupées par jour",
        "image_not_available": "Image non disponible",
        "map_not_available": "Carte non disponible",
        "day_prefix": "Jour",
    },
}


def _get_docx_labels(language: str) -> Dict[str, str]:
    """Get language-specific labels for DOCX generation."""
    return DOCX_LABELS.get(language, DOCX_LABELS["en"])


def add_horizontal_line(paragraph, color: RGBColor = COLORS["primary"], width: float = 1.0):
    """Add a horizontal line below a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(width * 8)))  # Size in eighths of a point
    bottom.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


class LocalDocxGenerator:
    """
    Local DOCX document generator using python-docx with modern styling.
    """

    def __init__(self, output_dir: str = "./.results"):
        """
        Initialize local DOCX generator.

        Args:
            output_dir: Directory to save generated documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        LOGGER.info(f"LocalDocxGenerator initialized with output_dir: {output_dir}")

    def _setup_document_styles(self, doc: Document):
        """Configure modern styles for the document."""
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.font.color.rgb = COLORS["text"]
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15

        # Configure Heading 1 (Day headers)
        heading1 = doc.styles['Heading 1']
        heading1.font.name = 'Calibri Light'
        heading1.font.size = Pt(24)
        heading1.font.color.rgb = COLORS["primary"]
        heading1.font.bold = False
        heading1.paragraph_format.space_before = Pt(24)
        heading1.paragraph_format.space_after = Pt(12)

        # Configure Heading 2 (Attraction names)
        heading2 = doc.styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(16)
        heading2.font.color.rgb = COLORS["dark"]
        heading2.font.bold = True
        heading2.paragraph_format.space_before = Pt(16)
        heading2.paragraph_format.space_after = Pt(8)

        # Configure Heading 3 (Subsections)
        heading3 = doc.styles['Heading 3']
        heading3.font.name = 'Calibri'
        heading3.font.size = Pt(12)
        heading3.font.color.rgb = COLORS["secondary"]
        heading3.font.bold = True
        heading3.paragraph_format.space_before = Pt(12)
        heading3.paragraph_format.space_after = Pt(4)


    def _add_styled_title(self, doc: Document, title: str, labels: Dict[str, str]):
        """Add a modern styled title with decorative line."""
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add title text
        run = title_para.add_run(title)
        run.font.name = 'Calibri Light'
        run.font.size = Pt(36)
        run.font.color.rgb = COLORS["primary"]

        # Add subtitle line
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"✈ {labels['travel_itinerary']} ✈")
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = COLORS["light_text"]
        run.font.italic = True

        # Add decorative line
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line_para.add_run("─" * 40)
        run.font.color.rgb = COLORS["secondary"]

        # Add spacing
        doc.add_paragraph()

    def _add_day_header(self, doc: Document, day_number: int, day_label: str = "Day"):
        """Add a modern styled day header with icon."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = para.add_run(f"{day_label} {day_number}")
        run.font.name = 'Calibri Light'
        run.font.size = Pt(26)
        run.font.color.rgb = COLORS["primary"]

        # Add underline
        add_horizontal_line(para, COLORS["secondary"], 2.0)

        para.paragraph_format.space_before = Pt(28)
        para.paragraph_format.space_after = Pt(16)

    def _add_attraction_header(self, doc: Document, name: str):
        """Add a styled attraction header."""
        para = doc.add_paragraph()

        # Attraction name
        run = para.add_run(name)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.color.rgb = COLORS["dark"]
        run.font.bold = True

        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)

    def _add_styled_bullet_list(self, doc: Document, items: List[str]):
        """Add a styled bullet list with custom formatting."""
        for item in items:
            para = doc.add_paragraph()
            # Bullet point icon
            run = para.add_run("• ")
            run.font.color.rgb = COLORS["secondary"]
            run.font.size = Pt(11)

            # Item text
            run = para.add_run(item)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = COLORS["text"]

            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.space_after = Pt(4)

    def _add_info_section(self, doc: Document, section_title: str):
        """Add a styled info section header."""
        para = doc.add_paragraph()

        run = para.add_run(section_title)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = COLORS["secondary"]
        run.font.bold = True

        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def _add_map_header(self, doc: Document, labels: Dict[str, str]):
        """Add a styled map section header."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run(labels["route_map"])
        run.font.name = 'Calibri Light'
        run.font.size = Pt(24)
        run.font.color.rgb = COLORS["primary"]

        para.paragraph_format.space_before = Pt(16)
        para.paragraph_format.space_after = Pt(16)

    def _detect_day_prefix(self, text: str) -> tuple[str, int]:
        """
        Detect the day prefix and number from a heading text.
        Returns (day_prefix, day_number) or (None, 0) if not a day header.
        """
        for lang_labels in DOCX_LABELS.values():
            prefix = lang_labels["day_prefix"]
            if text.startswith(f"{prefix} "):
                try:
                    day_num = int(text.split()[1])
                    return prefix, day_num
                except (ValueError, IndexError):
                    return prefix, 0
        return None, 0

    def create_document(
        self,
        title: str,
        content_blocks: List[Dict[str, Any]],
        output_filename: Optional[str] = None,
        language: str = "en"
    ) -> str:
        """
        Create a DOCX document with formatted content and images.

        Args:
            title: Document title
            content_blocks: List of content blocks (text, images, etc.)
            output_filename: Optional custom filename
            language: Language code for localized strings (en, pt-br, es, fr)

        Returns:
            Path to the created document
        """
        LOGGER.info("=== DOCX CREATION START ===")
        LOGGER.info(f"Title: {title}")
        LOGGER.info(f"Content blocks: {len(content_blocks)}")
        LOGGER.info(f"Language: {language}")

        # Get language-specific labels
        labels = _get_docx_labels(language)

        try:
            # Create document
            doc = Document()

            # Setup modern styles
            self._setup_document_styles(doc)

            # Add styled title
            LOGGER.info(f"Adding title: {title}")
            self._add_styled_title(doc, title, labels)

            # Process content blocks
            LOGGER.info("Processing content blocks...")
            for i, block in enumerate(content_blocks):
                block_type = block.get("type")
                LOGGER.debug(f"Block {i+1}/{len(content_blocks)}: type={block_type}")

                if block_type == "heading":
                    level = block.get("level", 1)
                    text = block.get("text", "")

                    if level == 1:
                        # Check if this is a day header
                        day_prefix, day_num = self._detect_day_prefix(text)
                        if day_prefix:
                            self._add_day_header(doc, day_num, day_prefix)
                        else:
                            heading = doc.add_heading(text, level)
                            heading.runs[0].font.color.rgb = COLORS["primary"]
                    elif level == 2:
                        self._add_attraction_header(doc, text)
                    elif level == 3:
                        self._add_info_section(doc, text)

                elif block_type == "paragraph":
                    text = block.get("text", "")
                    if not text:
                        doc.add_paragraph()
                        continue

                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.color.rgb = COLORS["text"]

                    # Apply formatting if specified
                    if block.get("bold"):
                        run.bold = True
                    if block.get("italic"):
                        run.italic = True

                elif block_type == "image":
                    image_url = block.get("url")
                    caption = block.get("caption", "")

                    LOGGER.info(f"Processing image: {image_url[:100]}...")

                    try:
                        # Download image
                        LOGGER.info(f"Downloading image from: {image_url}")
                        response = requests.get(image_url, timeout=30)

                        if response.status_code == 200:
                            LOGGER.info(f"Image downloaded successfully ({len(response.content)} bytes)")

                            # Load image to check dimensions
                            img = Image.open(BytesIO(response.content))
                            LOGGER.info(f"Image opened: {img.size}, format: {img.format}")

                            # Save to temp file
                            temp_path = os.path.join(self.output_dir, f"temp_{block.get('id', 'img')}.jpg")
                            LOGGER.info(f"Saving to temp file: {temp_path}")

                            img.save(temp_path, "JPEG")
                            LOGGER.info("Temp file saved successfully")

                            # Add to document (max width 5.5 inches for better margins)
                            LOGGER.info("Adding picture to document...")
                            doc.add_picture(temp_path, width=Inches(5.5))
                            LOGGER.info("Picture added successfully")

                            # Add styled caption
                            if caption:
                                caption_para = doc.add_paragraph()
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                                run = caption_para.add_run(caption)
                                run.font.name = 'Calibri'
                                run.font.size = Pt(9)
                                run.font.color.rgb = COLORS["light_text"]
                                run.italic = True

                                caption_para.paragraph_format.space_after = Pt(12)

                            # Clean up temp file
                            os.remove(temp_path)
                            LOGGER.info("Temp file removed")
                        else:
                            LOGGER.warning(f"Failed to download image: HTTP {response.status_code}")

                    except Exception as e:
                        LOGGER.error(f"Error adding image: {e}", exc_info=True)

                elif block_type == "bullet_list":
                    items = block.get("items", [])
                    self._add_styled_bullet_list(doc, items)

                elif block_type == "page_break":
                    doc.add_page_break()
                    LOGGER.debug("Page break added")

                elif block_type == "final_image":
                    doc.add_page_break()
                    self._add_map_header(doc, labels)

                    clusters = block.get("clusters", [])
                    attractions_coordinates = block.get("attraction_coordinates", {})
                    map_title = block.get("title", labels["route_map"])

                    if clusters.tolist() and attractions_coordinates:
                        attraction_names = list(attractions_coordinates.keys())
                        locs = [(attractions_coordinates[name]['lon'], attractions_coordinates[name]['lat']) for name in attraction_names]

                        map_image_path = os.path.join(self.output_dir, "final_map.png")
                        try:
                            plot_clusters_on_basemap(
                                locations=locs,
                                clusters=clusters,
                                names=attraction_names,
                                out_path=map_image_path,
                                title=map_title
                            )

                            LOGGER.info(f"Adding final map image: {map_image_path}")
                            doc.add_picture(map_image_path, width=Inches(6))
                            os.remove(map_image_path)

                            # Add legend
                            legend_para = doc.add_paragraph()
                            legend_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = legend_para.add_run(labels["map_legend"])
                            run.font.name = 'Calibri'
                            run.font.size = Pt(9)
                            run.font.color.rgb = COLORS["light_text"]
                            run.italic = True

                        except Exception as e:
                            LOGGER.error(f"Error generating final map image: {e}", exc_info=True)
                            para = doc.add_paragraph()
                            run = para.add_run(f"[{labels['map_not_available']}]")
                            run.font.color.rgb = COLORS["light_text"]
                    else:
                        LOGGER.warning("No clusters or coordinates provided for final image.")

            # Save document
            LOGGER.info("Saving document...")
            if not output_filename:
                import re
                safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
                output_filename = f"{safe_title}.docx"
                LOGGER.info(f"Generated filename: {output_filename}")

            output_path = os.path.join(self.output_dir, output_filename)
            LOGGER.info(f"Full output path: {output_path}")

            doc.save(output_path)
            LOGGER.info(f"Document saved successfully to: {output_path}")

            # Verify file exists
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                LOGGER.info(f"File verified: {file_size} bytes")
            else:
                LOGGER.error(f"File was not created at: {output_path}")

            LOGGER.info(f"Document saved: {output_path}")
            return output_path

        except Exception as e:
            LOGGER.error(f"Error creating DOCX: {e}", exc_info=True)
            return ""

    